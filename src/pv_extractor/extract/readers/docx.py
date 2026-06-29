"""DocxReader (D1): python-docx paragraphs + tables.

Word documents do not paginate without rendering, so the document is split
into pseudo-pages ("sections") at Heading 1/2 paragraphs — portfolio reviews
put each asset under its own heading, which is exactly the granularity
targeting (D2) and multi-asset row splitting need. A document without
headings is a single section.

Legacy .doc is NOT parsed: DocReaderUnsupported emits UNSUPPORTED_FORMAT so
the reviewer converts the file manually (spec D1).
"""

from __future__ import annotations

import io
from pathlib import Path

from pv_extractor.models import DocFlag, DocumentContent, PageContent, TableData

from .base import DocumentReader

_HEADING_STYLES = {"heading 1", "heading 2", "title"}


class DocxReader(DocumentReader):
    name = "docx"

    def summarize(self, path: str | Path, max_pages: int | None = None) -> DocumentContent:
        try:
            data = self._read_bytes(path)
        except OSError as exc:
            return self._empty(path, DocFlag.ACCESS_ERROR, f"{type(exc).__name__}: {exc}")
        if not data:
            return self._empty(path, DocFlag.CORRUPT_FILE, "zero-byte file")
        try:
            import docx

            document = docx.Document(io.BytesIO(data))
        except Exception as exc:
            return self._empty(path, DocFlag.CORRUPT_FILE, f"{type(exc).__name__}: {exc}")

        pages = _split_sections(document)
        shown = pages if max_pages is None else pages[:max_pages]
        return DocumentContent(
            file_path=str(path), reader=self.name, page_count=len(pages), pages=shown
        )

    def extract_tables(self, path: str | Path, page_numbers: list[int]) -> dict[int, list[TableData]]:
        content = self.summarize(path)
        wanted = set(page_numbers)
        return {
            page.page_number: page.tables
            for page in content.pages
            if page.page_number in wanted and page.tables
        }


class DocReaderUnsupported(DocumentReader):
    """Legacy binary .doc — never parsed; the reviewer converts manually."""

    name = "doc"

    def summarize(self, path: str | Path, max_pages: int | None = None) -> DocumentContent:
        return self._empty(
            path, DocFlag.UNSUPPORTED_FORMAT,
            "legacy .doc is not parsed; convert to .docx manually",
        )

    def extract_tables(self, path: str | Path, page_numbers: list[int]) -> dict[int, list[TableData]]:
        return {}


def _split_sections(document) -> list[PageContent]:
    """One PageContent per heading-delimited section, tables included
    (python-docx iterates body order via iter_inner_content)."""
    sections: list[PageContent] = []

    def new_section(name: str | None) -> PageContent:
        section = PageContent(
            page_number=len(sections) + 1, unit_label="section", unit_name=name
        )
        sections.append(section)
        return section

    current: PageContent | None = None
    texts: list[list[str]] = []

    import docx.table

    for block in document.iter_inner_content():
        if isinstance(block, docx.table.Table):
            if current is None:
                current = new_section(None)
                texts.append([])
            rows = [
                [cell.text.strip() or None for cell in row.cells] for row in block.rows
            ]
            if any(any(cell for cell in row) for row in rows):
                current.tables.append(
                    TableData(page_number=current.page_number, rows=rows, source="docx")
                )
                texts[-1].extend(
                    " ".join(str(cell) for cell in row if cell) for row in rows
                )
            continue
        # paragraph
        style = (block.style.name or "").lower() if block.style is not None else ""
        text = block.text
        if style in _HEADING_STYLES and text.strip() and (
            current is None or texts[-1]
        ):
            current = new_section(text.strip())
            texts.append([])
        elif current is None:
            current = new_section(None)
            texts.append([])
        if text:
            texts[-1].append(text)

    if current is None:
        new_section(None)
        texts.append([])

    for section, lines in zip(sections, texts):
        section.text = "\n".join(lines)
        section.text_char_count = len(section.text.strip())
        section.has_text_layer = True
    return sections
